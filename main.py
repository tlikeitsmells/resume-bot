import os
from telegram import Update, InputFile, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import Application, CommandHandler, MessageHandler, ConversationHandler, filters, ContextTypes
from docx import Document
from storage import Store
from tailoring import job_keywords, score_bullet, suggest_rewrite
from ai_helper import suggest_skills, suggest_bullets, refine_with_openai

# ===== States for /build guided interview =====
(CONTACT_NAME, CONTACT_EMAIL, CONTACT_PHONE, CONTACT_LOCATION, CONTACT_LINKS,
 EXP_ADD_ROLE_Q, EXP_TITLE, EXP_COMPANY, EXP_LOC, EXP_START, EXP_END, EXP_AI_BULLETS_Q, EXP_BULLET, EXP_ANOTHER_ROLE_Q,
 EDU_ADD_Q, EDU_DEGREE, EDU_SCHOOL, EDU_LOC, EDU_ENDYEAR, EDU_ANOTHER_Q,
 SKILLS_CORE, SKILLS_AI_Q, SKILLS_TOOLS, SKILLS_CERTS,
 SUMMARY_TEXT) = range(24)

store = Store()

def ats_docx(profile, path="resume_output.docx"):
    doc = Document()
    c = profile["contact"]
    doc.add_heading(c.get("name","").strip() or "Your Name", 0)
    line = " · ".join([x for x in [c.get("location",""), c.get("phone",""), c.get("email","")] if x])
    if line:
        doc.add_paragraph(line)
    if c.get("links"):
        doc.add_paragraph(" · ".join(c["links"]))

    if profile.get("summary"):
        doc.add_heading("Summary", 2); doc.add_paragraph(profile["summary"])

    skills = profile.get("skills", {})
    if any(skills.get(k) for k in ("core","tools","certs")):
        doc.add_heading("Skills", 2)
        chunks = []
        if skills.get("core"):  chunks.append(", ".join(skills["core"]))
        if skills.get("tools"): chunks.append(", ".join(skills["tools"]))
        if skills.get("certs"): chunks.append(", ".join(skills["certs"]))
        doc.add_paragraph("; ".join(chunks))

    if profile.get("experience"):
        doc.add_heading("Experience", 2)
        for r in profile["experience"]:
            h = f"{r['title']} — {r['company']} · {r['location']} ({r['start']}–{r['end']})"
            doc.add_paragraph(h)
            for b in r.get("bullets", []):
                doc.add_paragraph(f"• {b}")

    if profile.get("education"):
        doc.add_heading("Education", 2)
        for e in profile["education"]:
            p = f"{e['degree']} — {e['school']} · {e['location']} ({e['end']})"
            doc.add_paragraph(p)
            for h in e.get("highlights", []):
                doc.add_paragraph(f"• {h}")
    doc.save(path)
    return path

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Hey! I’ll build an ATS-friendly résumé with you.\\n"
        "Use /build for a guided interview, or commands: /tailor  /export  /show  /assist"
    )

async def export(update: Update, context: ContextTypes.DEFAULT_TYPE):
    p = store.get_profile(update.effective_user.id)
    path = ats_docx(p, "resume_output.docx")
    await update.message.reply_document(InputFile(path), filename="resume.docx")

async def show(update: Update, context: ContextTypes.DEFAULT_TYPE):
    p = store.get_profile(update.effective_user.id)
    lines = [
        f"*{p['contact'].get('name','')}*",
        p.get('summary',''),
        "",
        "Skills: " + ", ".join(p.get('skills',{}).get('core',[]))
    ]
    await update.message.reply_markdown_v2("\\n".join(lines))

async def assist(update: Update, context: ContextTypes.DEFAULT_TYPE):
    role = " ".join(context.args) if context.args else ""
    if not role:
        await update.message.reply_text("Usage: /assist <role>  (e.g., /assist Restaurant Manager)")
        return
    skills = suggest_skills(role)
    bullets = refine_with_openai(f"Rewrite bullets for '{role}' as outcome-first resume bullets.", suggest_bullets(role, 5))
    msg = "Suggested skills:\\n- " + "\\n- ".join(skills[:12]) + "\\n\\nSuggested bullets:\\n- " + "\\n- ".join(bullets)
    await update.message.reply_text(msg)

async def build_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["profile"] = store.get_profile(update.effective_user.id)
    context.user_data["profile"]["experience"] = []
    context.user_data["profile"]["education"] = []
    await update.message.reply_text("Let's build your résumé. First, your full name:")
    return CONTACT_NAME

async def contact_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["profile"]["contact"]["name"] = update.message.text.strip()
    await update.message.reply_text("Email address:")
    return CONTACT_EMAIL

async def contact_email(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["profile"]["contact"]["email"] = update.message.text.strip()
    await update.message.reply_text("Phone number:")
    return CONTACT_PHONE

async def contact_phone(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["profile"]["contact"]["phone"] = update.message.text.strip()
    await update.message.reply_text("City, ST (location):")
    return CONTACT_LOCATION

async def contact_location(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["profile"]["contact"]["location"] = update.message.text.strip()
    await update.message.reply_text("Any links? (comma-separated, or 'none')")
    return CONTACT_LINKS

async def contact_links(update: Update, context: ContextTypes.DEFAULT_TYPE):
    txt = update.message.text.strip()
    links = [] if txt.lower() == "none" else [x.strip() for x in txt.split(",") if x.strip()]
    context.user_data["profile"]["contact"]["links"] = links
    await update.message.reply_text("Add a job experience? (yes/no)",
        reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
    return EXP_ADD_ROLE_Q

async def exp_add_role_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "no":
        await update.message.reply_text("Add an education entry? (yes/no)",
            reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
        return EDU_ADD_Q
    await update.message.reply_text("Job Title:", reply_markup=ReplyKeyboardRemove())
    context.user_data["current_role"] = {"title":"","company":"","location":"","start":"","end":"","bullets":[]}
    return EXP_TITLE

async def exp_title(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_role"]["title"] = update.message.text.strip()
    await update.message.reply_text("Company:")
    return EXP_COMPANY

async def exp_company(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_role"]["company"] = update.message.text.strip()
    await update.message.reply_text("City, ST:")
    return EXP_LOC

async def exp_loc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_role"]["location"] = update.message.text.strip()
    await update.message.reply_text("Start date (YYYY-MM):")
    return EXP_START

async def exp_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_role"]["start"] = update.message.text.strip()
    await update.message.reply_text("End date (YYYY or 'Present'):")
    return EXP_END

async def exp_end(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_role"]["end"] = update.message.text.strip()
    role = context.user_data["current_role"]["title"]
    await update.message.reply_text(
        f"Want suggested bullets for '{role}'? (yes/no)",
        reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True)
    )
    return EXP_AI_BULLETS_Q

async def exp_ai_bullets_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "yes":
        role = context.user_data["current_role"]["title"]
        bullets = refine_with_openai(f"Rewrite as concise, outcome-first bullets for '{role}'.", suggest_bullets(role, 5))
        context.user_data["current_role"]["bullets"].extend(bullets)
        await update.message.reply_text("Added 5 suggested bullets. You can add more—send a bullet per message, or type /done when finished.",
            reply_markup=ReplyKeyboardRemove())
        return EXP_BULLET
    await update.message.reply_text("Add bullets one by one (send each as a message). Type /done when finished.",
        reply_markup=ReplyKeyboardRemove())
    return EXP_BULLET

async def exp_bullet(update: Update, context: ContextTypes.DEFAULT_TYPE):
    txt = update.message.text.strip()
    if txt.startswith("/"):
        return EXP_BULLET
    context.user_data["current_role"]["bullets"].append(txt)
    await update.message.reply_text("Added. Another bullet or /done when finished.")
    return EXP_BULLET

async def exp_done(update: Update, context: ContextTypes.DEFAULT_TYPE):
    role = context.user_data.get("current_role", {})
    context.user_data["profile"]["experience"].insert(0, role)
    context.user_data["current_role"] = {}
    await update.message.reply_text("Add another job? (yes/no)",
        reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
    return EXP_ANOTHER_ROLE_Q

async def exp_another_role_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "yes":
        await update.message.reply_text("Job Title:", reply_markup=ReplyKeyboardRemove())
        context.user_data["current_role"] = {"title":"","company":"","location":"","start":"","end":"","bullets":[]}
        return EXP_TITLE
    await update.message.reply_text("Add an education entry? (yes/no)",
        reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
    return EDU_ADD_Q

async def edu_add_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "no":
        role_hint = context.user_data["profile"]["experience"][0]["title"] if context.user_data["profile"]["experience"] else ""
        suggested = suggest_skills(role_hint)
        await update.message.reply_text("Core skills (comma-separated). Example:\\n" + ", ".join(suggested[:10]))
        return SKILLS_CORE
    await update.message.reply_text("Degree (e.g., BA Culinary Arts):", reply_markup=ReplyKeyboardRemove())
    context.user_data["current_edu"] = {"degree":"","school":"","location":"","end":"","highlights":[]}
    return EDU_DEGREE

async def edu_degree(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_edu"]["degree"] = update.message.text.strip()
    await update.message.reply_text("School:")
    return EDU_SCHOOL

async def edu_school(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_edu"]["school"] = update.message.text.strip()
    await update.message.reply_text("City, ST:")
    return EDU_LOC

async def edu_loc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_edu"]["location"] = update.message.text.strip()
    await update.message.reply_text("Graduation year (YYYY):")
    return EDU_ENDYEAR

async def edu_endyear(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data["current_edu"]["end"] = update.message.text.strip()
    context.user_data["profile"]["education"].append(context.user_data["current_edu"])
    context.user_data["current_edu"] = {}
    await update.message.reply_text("Add another education entry? (yes/no)",
        reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
    return EDU_ANOTHER_Q

async def edu_another_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "yes":
        await update.message.reply_text("Degree (e.g., BA Culinary Arts):", reply_markup=ReplyKeyboardRemove())
        context.user_data["current_edu"] = {"degree":"","school":"","location":"","end":"","highlights":[]}
        return EDU_DEGREE
    role_hint = context.user_data["profile"]["experience"][0]["title"] if context.user_data["profile"]["experience"] else ""
    suggested = suggest_skills(role_hint)
    await update.message.reply_text("Core skills (comma-separated). Example:\\n" + ", ".join(suggested[:10]), reply_markup=ReplyKeyboardRemove())
    return SKILLS_CORE

async def skills_core(update: Update, context: ContextTypes.DEFAULT_TYPE):
    core = [s.strip() for s in update.message.text.split(",") if s.strip()]
    context.user_data["profile"]["skills"]["core"] = core
    role_hint = context.user_data["profile"]["experience"][0]["title"] if context.user_data["profile"]["experience"] else ""
    suggested = [s for s in suggest_skills(role_hint) if s not in core][:8]
    if suggested:
        await update.message.reply_text("Add these suggested skills too? (yes/no)\\n" + ", ".join(suggested),
            reply_markup=ReplyKeyboardMarkup([['yes','no']], one_time_keyboard=True, resize_keyboard=True))
        context.user_data["suggested_skills"] = suggested
        return SKILLS_AI_Q
    await update.message.reply_text("Tools (comma-separated, or 'none'):")
    return SKILLS_TOOLS

async def skills_ai_q(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ans = update.message.text.strip().lower()
    if ans == "yes":
        context.user_data["profile"]["skills"]["core"].extend(context.user_data.get("suggested_skills", []))
    await update.message.reply_text("Tools (comma-separated, or 'none'):", reply_markup=ReplyKeyboardRemove())
    return SKILLS_TOOLS

async def skills_tools(update: Update, context: ContextTypes.DEFAULT_TYPE):
    txt = update.message.text.strip()
    tools = [] if txt.lower() == "none" else [s.strip() for s in txt.split(",") if s.strip()]
    context.user_data["profile"]["skills"]["tools"] = tools
    await update.message.reply_text("Certifications (comma-separated, or 'none'):")
    return SKILLS_CERTS

async def skills_certs(update: Update, context: ContextTypes.DEFAULT_TYPE):
    txt = update.message.text.strip()
    certs = [] if txt.lower() == "none" else [s.strip() for s in txt.split(",") if s.strip()]
    context.user_data["profile"]["skills"]["certs"] = certs
    await update.message.reply_text("Paste a tight 2–3 line professional summary (or 'skip'):")
    return SUMMARY_TEXT

async def summary_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    txt = update.message.text.strip()
    if txt.lower() != "skip":
        context.user_data["profile"]["summary"] = txt
    store.set_profile(update.effective_user.id, context.user_data["profile"])
    await update.message.reply_text("Saved! Use /export to download your DOCX, or /tailor to paste a job description.")
    return ConversationHandler.END

async def tailor(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Paste the job description. I’ll rank and suggest bullet tweaks.")
    return 1000

async def tailor_capture(update: Update, context: ContextTypes.DEFAULT_TYPE):
    jd = update.message.text
    kw = job_keywords(jd)
    p = store.get_profile(update.effective_user.id)
    ranked = []
    for r in p["experience"]:
        for b in r.get("bullets", []):
            ranked.append((score_bullet(b, kw), b, r["title"]))
    ranked.sort(reverse=True, key=lambda x: x[0])
    top = ranked[:8]
    if not top:
        await update.message.reply_text("No bullets yet. Add experience via /build.")
        return ConversationHandler.END
    suggestions = [f"• {suggest_rewrite(b, kw)}  (from {t})" for _, b, t in top]
    await update.message.reply_text("\\n".join(suggestions))
    return ConversationHandler.END

def build_application():
    app = Application.builder().token(os.environ["TELEGRAM_BOT_TOKEN"]).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("show", show))
    app.add_handler(CommandHandler("export", export))
    app.add_handler(CommandHandler("assist", assist))

    build_conv = ConversationHandler(
        entry_points=[CommandHandler("build", build_start)],
        states={
            CONTACT_NAME:   [MessageHandler(filters.TEXT & ~filters.COMMAND, contact_name)],
            CONTACT_EMAIL:  [MessageHandler(filters.TEXT & ~filters.COMMAND, contact_email)],
            CONTACT_PHONE:  [MessageHandler(filters.TEXT & ~filters.COMMAND, contact_phone)],
            CONTACT_LOCATION:[MessageHandler(filters.TEXT & ~filters.COMMAND, contact_location)],
            CONTACT_LINKS:  [MessageHandler(filters.TEXT & ~filters.COMMAND, contact_links)],

            EXP_ADD_ROLE_Q: [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_add_role_q)],
            EXP_TITLE:      [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_title)],
            EXP_COMPANY:    [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_company)],
            EXP_LOC:        [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_loc)],
            EXP_START:      [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_start)],
            EXP_END:        [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_end)],
            EXP_AI_BULLETS_Q: [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_ai_bullets_q)],
            EXP_BULLET:     [
                MessageHandler(filters.TEXT & ~filters.COMMAND, exp_bullet),
                CommandHandler("done", exp_done)
            ],
            EXP_ANOTHER_ROLE_Q: [MessageHandler(filters.TEXT & ~filters.COMMAND, exp_another_role_q)],

            EDU_ADD_Q:      [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_add_q)],
            EDU_DEGREE:     [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_degree)],
            EDU_SCHOOL:     [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_school)],
            EDU_LOC:        [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_loc)],
            EDU_ENDYEAR:    [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_endyear)],
            EDU_ANOTHER_Q:  [MessageHandler(filters.TEXT & ~filters.COMMAND, edu_another_q)],

            SKILLS_CORE:    [MessageHandler(filters.TEXT & ~filters.COMMAND, skills_core)],
            SKILLS_AI_Q:    [MessageHandler(filters.TEXT & ~filters.COMMAND, skills_ai_q)],
            SKILLS_TOOLS:   [MessageHandler(filters.TEXT & ~filters.COMMAND, skills_tools)],
            SKILLS_CERTS:   [MessageHandler(filters.TEXT & ~filters.COMMAND, skills_certs)],
            SUMMARY_TEXT:   [MessageHandler(filters.TEXT & ~filters.COMMAND, summary_text)],
        },
        fallbacks=[CommandHandler("cancel", lambda u,c: ConversationHandler.END)]
    )
    app.add_handler(build_conv)

    tailor_conv = ConversationHandler(
        entry_points=[CommandHandler("tailor", tailor)],
        states={1000:[MessageHandler(filters.TEXT & ~filters.COMMAND, tailor_capture)]},
        fallbacks=[CommandHandler("cancel", lambda u,c: ConversationHandler.END)]
    )
    app.add_handler(tailor_conv)

    return app

if __name__ == "__main__":
    application = build_application()
    application.run_polling()
